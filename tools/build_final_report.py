from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = Path(r"C:\Users\bjk-6\Downloads")
OUT = ROOT / "Med-AgentLab_Bitirme_Projesi_Final_Rapor.docx"
LOGO = DOWNLOADS / "samsunün.jpg"
ASSET_DIR = ROOT / "report_assets"
LOGO_PNG = ASSET_DIR / "samsun_logo.png"


TITLE = "Med-AgentLab: Tıp Dikeyinde Gizlilik Odaklı, Kota Duyarlı Çoklu Ajan Orkestrasyon Altyapısı"
SHORT_TITLE = "Med-AgentLab"
STUDENT = "Emin ALTAN"
DEPARTMENT = "Yazılım Mühendisliği Bölümü"
FACULTY = "Mühendislik ve Doğa Bilimleri Fakültesi"
UNIVERSITY = "Samsun Üniversitesi"
SUPERVISOR = "[Danışman Unvanı Adı SOYADI]"
MONTH_YEAR = "Haziran 2026"


def tr_upper(text: str) -> str:
    table = str.maketrans({"i": "İ", "ı": "I", "ğ": "Ğ", "ü": "Ü", "ş": "Ş", "ö": "Ö", "ç": "Ç"})
    return text.translate(table).upper()


def set_cell_text(cell, text: str, bold: bool = False, size: int = 11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: Sequence[float]):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
                set_cell_margins(row.cells[idx])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def apply_document_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(6)
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6)


def add_page_number(section, fmt: str = "decimal", start: int = 1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), fmt)

    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for child in list(p._p):
        p._p.remove(child)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = str(start)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def break_page(doc: Document):
    doc.add_page_break()


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False, after: int = 6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    return p


def add_right(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    return p


def add_body_paragraph(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
    else:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(12)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(12)


def add_numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(12)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_cm: Sequence[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths_cm)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_text(hdr.cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=10 if len(value) > 120 else 11)
    set_table_width(table, widths_cm)
    doc.add_paragraph()
    return table


def add_caption(doc: Document, text: str, above: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    if "." in text:
        prefix = text.split(".", 1)[0] + "."
        p.clear()
        r1 = p.add_run(prefix)
        r1.bold = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(10)
        r2 = p.add_run(text[len(prefix):])
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)
    return p


def create_diagrams():
    ASSET_DIR.mkdir(exist_ok=True)
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return []

    if LOGO.exists():
        try:
            Image.open(LOGO).convert("RGB").save(LOGO_PNG)
        except Exception:
            pass

    font_path = "C:/Windows/Fonts/arial.ttf"
    bold_path = "C:/Windows/Fonts/arialbd.ttf"
    font = ImageFont.truetype(font_path, 24)
    small = ImageFont.truetype(font_path, 18)
    bold = ImageFont.truetype(bold_path, 24)

    def box(draw, xy, title, body, fill="#F4F8FB"):
        draw.rounded_rectangle(xy, radius=18, fill=fill, outline="#3B6EA5", width=3)
        x1, y1, x2, y2 = xy
        draw.text((x1 + 18, y1 + 16), title, fill="#0B2545", font=bold)
        for idx, line in enumerate(body):
            draw.text((x1 + 18, y1 + 52 + idx * 24), line, fill="#1E293B", font=small)

    img = Image.new("RGB", (1600, 780), "white")
    d = ImageDraw.Draw(img)
    boxes = [
        ((40, 90, 310, 260), "Web Arayüzü", ["Excel yükleme", "Canlı durum", "Sonuç/rapor"]),
        ((380, 90, 650, 260), "FastAPI", ["Upload endpoint", "BackgroundTasks", "Job durumu"]),
        ((720, 90, 990, 260), "Pipeline", ["Chunking", "Semaphore", "Ajan sırası"]),
        ((1060, 90, 1330, 260), "Çıktılar", ["Excel", "Markdown", "DOCX"]),
        ((210, 420, 500, 610), "Yerel Katman", ["Ollama qwen3:4b", "Regex PatternGuard", "PII redaksiyon"]),
        ((580, 420, 870, 610), "Model Havuzu", ["Groq", "Gemini", "Fallback router"]),
        ((950, 420, 1240, 610), "Doğrulama", ["PubMed E-utilities", "LLM denetimi", "PARTIAL/YES/NO"]),
    ]
    for b in boxes:
        box(d, *b)
    arrows = [((310, 175), (380, 175)), ((650, 175), (720, 175)), ((990, 175), (1060, 175)),
              ((855, 260), (355, 420)), ((855, 260), (725, 420)), ((855, 260), (1095, 420))]
    for a, b in arrows:
        d.line([a, b], fill="#B91C1C", width=5)
        d.ellipse((b[0]-7, b[1]-7, b[0]+7, b[1]+7), fill="#B91C1C")
    d.text((40, 24), "Med-AgentLab Sistem Mimarisi", fill="#B91C1C", font=ImageFont.truetype(bold_path, 34))
    arch_path = ASSET_DIR / "architecture.png"
    img.save(arch_path)

    img2 = Image.new("RGB", (1600, 520), "white")
    d = ImageDraw.Draw(img2)
    labels = [
        ("1. Girdi", "Excel/TXT/PDF"),
        ("2. Gizlilik", "Ollama + Regex"),
        ("3. Tema", "Model havuzu"),
        ("4. Validasyon", "PubMed + LLM"),
        ("5. Sentez", "Reducer"),
        ("6. Rapor", "XLSX/MD/DOCX"),
    ]
    x = 40
    for i, (t, b) in enumerate(labels):
        box(d, (x, 160, x + 220, 310), t, [b], fill="#FFF7ED" if i == 1 else "#F8FAFC")
        if i < len(labels) - 1:
            d.line([(x + 220, 235), (x + 285, 235)], fill="#0F766E", width=5)
            d.polygon([(x + 285, 235), (x + 265, 223), (x + 265, 247)], fill="#0F766E")
        x += 270
    d.text((40, 60), "Beş Aşamalı Analiz Akışı", fill="#0F766E", font=ImageFont.truetype(bold_path, 34))
    flow_path = ASSET_DIR / "pipeline_flow.png"
    img2.save(flow_path)
    return [arch_path, flow_path]


def add_cover(doc: Document):
    add_centered(doc, tr_upper(UNIVERSITY), 14, True, 3)
    add_centered(doc, tr_upper(FACULTY), 13, True, 18)
    logo_path = LOGO_PNG if LOGO_PNG.exists() else LOGO
    if logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run().add_picture(str(logo_path), width=Inches(1.75))
    for _ in range(3):
        doc.add_paragraph()
    add_centered(doc, tr_upper(TITLE), 14, True, 18)
    add_centered(doc, tr_upper(STUDENT), 12, True, 8)
    add_centered(doc, "BİTİRME PROJESİ", 12, True, 8)
    add_centered(doc, DEPARTMENT, 12, False, 8)
    add_centered(doc, f"Danışman: {SUPERVISOR}", 12, False, 36)
    add_centered(doc, MONTH_YEAR, 12, False, 0)


def add_inner_cover(doc: Document):
    add_centered(doc, tr_upper(UNIVERSITY), 14, True, 3)
    add_centered(doc, tr_upper(FACULTY), 13, True, 24)
    add_centered(doc, tr_upper(TITLE), 14, True, 22)
    add_centered(doc, tr_upper(STUDENT), 12, True, 8)
    add_centered(doc, "BİTİRME PROJESİ", 12, True, 8)
    add_centered(doc, DEPARTMENT, 12, False, 8)
    add_centered(doc, f"Danışman: {SUPERVISOR}", 12, False, 18)
    add_body_paragraph(doc, "Bu raporda kullanılan proje bilgileri; proje klasöründeki kaynak kod dosyaları, README/ABOUT/idea/Knowledge-Base metinleri ve yazım kılavuzu incelenerek hazırlanmıştır. Hasta verisi içeren demo veri dosyalarının içerikleri rapora aktarılmamıştır.")
    add_centered(doc, MONTH_YEAR, 12, False, 0)


def add_approval(doc: Document):
    add_centered(doc, "BİTİRME PROJESİ ONAYI", 14, True, 18)
    add_body_paragraph(doc, f"{STUDENT} tarafından hazırlanan “{TITLE}” başlıklı bitirme projesi, .../.../2026 tarihinde aşağıdaki jüri tarafından değerlendirilerek Samsun Üniversitesi {FACULTY} {DEPARTMENT} kapsamında Bitirme Projesi olarak kabul edilmiştir.")
    for label in ["Danışman", "Jüri Üyesi", "Jüri Üyesi", "Bölüm Başkanı"]:
        doc.add_paragraph()
        p = add_centered(doc, "........................................................", 12, False, 0)
        p.paragraph_format.space_before = Pt(12)
        add_centered(doc, f"({label} - Unvanı, Adı ve SOYADI)", 11, False, 8)


def add_declaration(doc: Document):
    add_right(doc, "..../..../2026")
    add_centered(doc, "ETİK İLKE VE KURALLARA UYGUNLUK BEYANNAMESİ", 14, True, 18)
    add_body_paragraph(doc, "Bu bitirme projesinin bana ait, özgün bir çalışma olduğunu; çalışmamın hazırlık, veri toplama, analiz ve bilgilerin sunumu olmak üzere tüm aşamalarında bilimsel etik ilke ve kurallara uygun davrandığımı; bu çalışma kapsamında yararlandığım tüm kaynaklara metin içinde atıf yaptığımı ve kaynakçada yer verdiğimi beyan ederim.")
    add_body_paragraph(doc, "Bu raporda proje kodu üzerinden doğrulanamayan özellik iddialarına yer verilmemiştir. Tıbbi örnek veri seti yalnızca sistemin test senaryosunu açıklamak için genel düzeyde ele alınmış; hasta görüşmesi metinleri, kişisel bilgi içeren ham örnekler veya bireysel hasta anlatıları rapora aktarılmamıştır.")
    for _ in range(3):
        doc.add_paragraph()
    add_right(doc, STUDENT)


def add_front_matter(doc: Document):
    add_heading(doc, "Önsöz", 1)
    add_body_paragraph(doc, "Bu bitirme projesi, nitel tıbbi veri analizinde maliyet, gizlilik ve ölçeklenebilirlik sorunlarına mühendislik odaklı bir çözüm geliştirme amacıyla hazırlanmıştır. Proje sürecinde tekil bir büyük dil modeli kullanımının yerine, görev bazlı ve kota duyarlı bir model havuzu yaklaşımı benimsenmiştir.")
    add_body_paragraph(doc, "Çalışmanın temel motivasyonu, kısıtlı donanıma sahip bir geliştirme ortamında yerel model avantajını korurken bulut tabanlı modellerin hız ve kapasite avantajlarından yararlanmaktır. Bu nedenle yerel Ollama katmanı kişisel veri azaltma göreviyle sınırlandırılmış, daha ağır tematik analiz ve rapor sentezi görevleri model havuzu üzerinden yürütülmüştür.")
    add_right(doc, STUDENT)
    add_right(doc, MONTH_YEAR)
    break_page(doc)

    add_heading(doc, "İçindekiler", 1)
    toc_items = [
        "Önsöz",
        "Özet",
        "Abstract",
        "Tablolar Listesi",
        "Şekiller Listesi",
        "Simgeler ve Kısaltmalar Dizini",
        "Giriş",
        "Birinci Bölüm: Kavramsal Çerçeve ve Problem Tanımı",
        "İkinci Bölüm: Sistem Tasarımı ve Mimari",
        "Üçüncü Bölüm: Uygulama, Deneysel Çalışma ve Değerlendirme",
        "Sonuç ve Öneriler",
        "Kaynakça",
        "Ekler",
        "Özgeçmiş",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(item).font.name = "Times New Roman"
    break_page(doc)

    add_heading(doc, "Özet", 1)
    add_centered(doc, tr_upper(TITLE), 12, True, 4)
    add_centered(doc, STUDENT, 12, True, 4)
    add_centered(doc, f"{DEPARTMENT}", 12, False, 4)
    add_centered(doc, f"{UNIVERSITY}, {FACULTY}, {MONTH_YEAR}", 12, False, 4)
    add_centered(doc, f"Danışman: {SUPERVISOR}", 12, False, 10)
    add_body_paragraph(doc, "Bu çalışmada tıp dikeyindeki nitel veri analiz süreçleri için gizlilik odaklı, maliyet-etkin ve çoklu model kullanabilen bir yazılım prototipi geliştirilmiştir. Med-AgentLab adlı sistem, Excel tabanlı klinik görüşme kayıtlarını alarak yerel gizlilik ön işleminden geçirir, tema haritalama aşamasında model havuzu kullanır, çıkarılan temaları PubMed destekli doğrulama yaklaşımıyla denetler ve sonuçları Excel, Markdown ve Word raporu olarak sunar. Sistem FastAPI tabanlı bir arka uç, tek sayfalık web arayüzü, SQLite iş geçmişi, LiteLLM üzerinden model çağrıları, Ollama tabanlı yerel model katmanı ve kural tabanlı fallback mekanizmalarından oluşmaktadır. Kod incelemesine göre sistemde kota duyarlı model yönlendirme, analiz iptali, canlı log takibi, rapor indirme ve demo-safe çalışma modu uygulanmıştır. Çalışma, gerçek klinik karar verme sistemi olarak değil, tıbbi nitel veri analizinde çoklu ajan orkestrasyonunun uygulanabilirliğini gösteren bir mühendislik prototipi olarak değerlendirilmelidir.")
    add_body_paragraph(doc, "Anahtar Sözcükler: Çoklu Ajan Sistemleri, Nitel Veri Analizi, MapReduce, Büyük Dil Modelleri, Gizlilik, LiteLLM.")
    break_page(doc)

    add_heading(doc, "Abstract", 1)
    add_centered(doc, tr_upper(TITLE), 12, True, 4)
    add_centered(doc, STUDENT, 12, True, 4)
    add_centered(doc, f"{DEPARTMENT}", 12, False, 4)
    add_centered(doc, f"{UNIVERSITY}, {FACULTY}, {MONTH_YEAR}", 12, False, 4)
    add_centered(doc, f"Supervisor: {SUPERVISOR}", 12, False, 10)
    add_body_paragraph(doc, "This graduation project presents Med-AgentLab, a privacy-oriented and cost-aware prototype for qualitative medical data analysis. The system processes Excel-based clinical interview records, applies a local privacy preprocessing layer, extracts themes through a quota-aware model pool, validates extracted themes with a PubMed-supported review stage, and exports results as Excel, Markdown, and Word reports. The implementation consists of a FastAPI backend, a single-page web interface, SQLite-based job history, LiteLLM-based model calls, a local Ollama layer, and deterministic fallback mechanisms. Based on source-code inspection, the implemented system supports background analysis jobs, live progress logs, job cancellation, model router events, result visualization, and report download endpoints. The project should be considered an engineering prototype for multi-agent orchestration in qualitative medical analysis rather than a clinically validated decision-support system.")
    add_body_paragraph(doc, "Keywords: Multi-Agent Systems, Qualitative Data Analysis, MapReduce, Large Language Models, Privacy, LiteLLM.")
    break_page(doc)

    add_heading(doc, "Tablolar Listesi", 1)
    for item in [
        "TABLO 1. Kullanılan Teknolojiler ve Projedeki Görevleri",
        "TABLO 2. Uygulanan Analiz Boru Hattı Aşamaları",
        "TABLO 3. API Uç Noktaları ve İşlevleri",
        "TABLO 4. Sınırlılıklar ve Risk Azaltma Yaklaşımları",
    ]:
        add_body_paragraph(doc, item)
    break_page(doc)

    add_heading(doc, "Şekiller Listesi", 1)
    add_body_paragraph(doc, "ŞEKİL 1. Med-AgentLab Sistem Mimarisi")
    add_body_paragraph(doc, "ŞEKİL 2. Beş Aşamalı Analiz Akışı")
    break_page(doc)

    add_heading(doc, "Simgeler ve Kısaltmalar Dizini", 1)
    abbreviations = [
        ("API", "Application Programming Interface"),
        ("BERT/LLM", "Büyük Dil Modeli / Large Language Model"),
        ("DOCX", "Microsoft Word belge biçimi"),
        ("JSON", "JavaScript Object Notation"),
        ("KVKK", "Kişisel Verilerin Korunması Kanunu"),
        ("LLM", "Large Language Model"),
        ("PII", "Personally Identifiable Information"),
        ("RAG", "Retrieval-Augmented Generation"),
        ("SQLite", "Gömülü ilişkisel veritabanı"),
    ]
    add_table(doc, ["Kısaltma", "Açıklama"], abbreviations, [3.0, 11.5])
    break_page(doc)


def add_main_text(doc: Document, diagrams: Sequence[Path]):
    add_heading(doc, "Giriş", 1)
    add_body_paragraph(doc, "Tıbbi araştırmalarda hasta görüşmeleri, açık uçlu anketler ve klinik notlar gibi metinsel veriler sıklıkla nitel analiz yöntemleriyle incelenmektedir. Bu tür analizlerde araştırmacının tekrar eden anlam birimlerini, semptom örüntülerini ve deneyim temalarını sistematik biçimde kodlaması gerekir. Geleneksel nitel analiz araçları araştırmacıya düzenleme ve etiketleme kolaylığı sağlasa da, geniş hacimli metinlerin elle incelenmesi zaman ve maliyet açısından önemli bir darboğaz oluşturur.")
    add_body_paragraph(doc, "Büyük dil modelleri bu darboğazı azaltma potansiyeline sahip olmakla birlikte, tek bir modele uzun ve hassas klinik metinleri doğrudan göndermek bağlam penceresi, maliyet, kota ve gizlilik sorunlarını beraberinde getirir. Bu çalışmanın çıkış noktası, nitel analiz sürecini tekil model kullanımından çıkararak görev bazlı, denetlenebilir ve hata toleranslı bir orkestrasyon mimarisine dönüştürmektir.")
    add_body_paragraph(doc, "Bu kapsamda geliştirilen Med-AgentLab prototipi, veri alma, gizlilik ön işleme, tema haritalama, tıbbi doğrulama ve akademik sentez aşamalarından oluşan beş aşamalı bir boru hattı sunar. Sistem, yerel Ollama modelini hafif gizlilik görevi için kullanırken, daha ağır tema çıkarma ve rapor üretme görevlerinde kota duyarlı model havuzuna başvurur.")

    add_heading(doc, "Çalışmanın Amacı", 2)
    add_body_paragraph(doc, "Çalışmanın amacı, tıp dikeyindeki nitel metinlerin gizlilik kaygısı gözetilerek otomatik biçimde işlenmesini sağlayan, çoklu model kullanabilen ve ücretsiz/kısıtlı API koşullarında hata toleransı gösterebilen bir yazılım prototipi geliştirmektir.")

    add_heading(doc, "Çalışmanın Kapsamı", 2)
    add_body_paragraph(doc, "Bu rapor, Med-AgentLab uygulamasının mevcut kod tabanı üzerinden doğrulanabilen özelliklerini kapsamaktadır. Rapor kapsamında hasta verisi içeren demo kayıtlarının ham içeriği kullanılmamış, yalnızca sistemin dosya işleme biçimi ve çıktı üretme kabiliyetleri değerlendirilmiştir.")

    add_heading(doc, "Araştırma ve Mühendislik Soruları", 2)
    add_numbered(doc, [
        "Kısıtlı yerel donanıma sahip bir ortamda yerel model ve bulut modelleri aynı analiz boru hattında nasıl konumlandırılabilir?",
        "API kota hataları ve model erişim sorunları analiz sürecini tamamen durdurmadan nasıl yönetilebilir?",
        "Nitel veri analizinde tema çıkarma, doğrulama ve rapor sentezi aşamaları birbirinden ayrılmış görevler olarak nasıl modellenebilir?",
        "Kullanıcıya analiz ilerlemesi, model yönlendirme ve çıktı indirme süreçleri nasıl izlenebilir biçimde sunulabilir?",
    ])

    add_heading(doc, "Birinci Bölüm", 1)
    add_heading(doc, "Kavramsal Çerçeve ve Problem Tanımı", 2)
    add_body_paragraph(doc, "Nitel veri analizi, metin içinde tekrar eden anlam örüntülerinin kodlanması, gruplanması ve yorumlanmasına dayanır. Braun ve Clarke tarafından sistematikleştirilen tematik analiz yaklaşımı, veriye aşinalık kazanma, kodlama, tema arama, temaları gözden geçirme ve raporlama gibi aşamalardan oluşur (Braun ve Clarke, 2006). Med-AgentLab bu metodolojik çerçeveyi yazılım düzeyinde doğrudan bire bir otomatikleştirmekten çok, kodlama ve tema sentezi görevlerini destekleyen bir mühendislik altyapısı olarak ele alır.")
    add_body_paragraph(doc, "MapReduce yaklaşımı, büyük veri kümelerinde map aşamasında parçalı işlem, reduce aşamasında ise ara çıktıların birleştirilmesi fikrine dayanır (Dean ve Ghemawat, 2004). Bu projede aynı kavram, klinik metin segmentleri üzerinde tema çıkarma ve doğrulanmış temaları rapor sentezinde birleştirme biçiminde uyarlanmıştır.")
    add_body_paragraph(doc, "Problem üç eksende ele alınmıştır: bağlam yönetimi, gizlilik ve maliyet/kota yönetimi. Bağlam yönetimi için uzun metinlerin segmentlere ayrılması; gizlilik için yerel Ollama ve regex tabanlı ön işleme; maliyet/kota yönetimi için ise görev bazlı model havuzu uygulanmıştır.")

    add_heading(doc, "Mevcut Yaklaşımların Sınırlılıkları", 2)
    add_body_paragraph(doc, "Geleneksel nitel analiz yazılımları araştırmacının manuel kodlama sürecini destekler; ancak analizi doğrudan otonom biçimde yürütmez. Tekil LLM tabanlı yaklaşımlar ise hızlı sonuç üretebilse de uzun metinlerde bağlam kaybı, hassas verinin buluta çıkması ve tek sağlayıcıya bağımlılık gibi riskler taşır.")
    add_body_paragraph(doc, "Med-AgentLab, bu riskleri tekil bir sohbet modelinin ötesine geçerek aşamalı ve görev odaklı bir sistem tasarımıyla azaltmaya çalışır. Bu yaklaşımda her model her işi yapmaz; yerel model dar kapsamlı gizlilik görevine, bulut modelleri ise daha ağır çıkarım görevlerine yönlendirilir.")

    add_heading(doc, "İkinci Bölüm", 1)
    add_heading(doc, "Sistem Tasarımı ve Mimari", 2)
    if diagrams:
        doc.add_picture(str(diagrams[0]), width=Inches(6.0))
        add_caption(doc, "Şekil 1. Med-AgentLab Sistem Mimarisi")
    add_body_paragraph(doc, "Uygulamanın güncel web sürümü FastAPI tabanlı app.py dosyasında yer almaktadır. Sistem; dosya yükleme, arka plan analiz işi başlatma, iş durumunu sorgulama, sonuçları JSON olarak döndürme, Excel/Markdown/DOCX indirme ve frontend dosyasını sunma uç noktalarından oluşur.")
    add_body_paragraph(doc, "Frontend tek sayfalık bir HTML uygulaması olarak frontend/index.html dosyasında geliştirilmiştir. Arayüzde dosya yükleme alanı, geçmiş analizler listesi, canlı terminal paneli, pipeline adımları, sonuç tablosu, tema grafiği, router ve gizlilik paneli ile rapor görüntüleme sekmeleri bulunmaktadır.")

    add_caption(doc, "Tablo 1. Kullanılan Teknolojiler ve Projedeki Görevleri")
    add_table(doc, ["Teknoloji", "Projede Kullanımı"], [
        ("Python", "Asenkron analiz boru hattı, veri işleme ve backend uygulaması."),
        ("FastAPI", "Upload, status, result, download ve cancel uç noktalarının sunulması."),
        ("LiteLLM", "Groq, Gemini ve Ollama gibi farklı model sağlayıcılarına ortak arayüz üzerinden çağrı yapılması."),
        ("Ollama", "Yerel gizlilik ve ön işleme ajanı için qwen3:4b modelinin çalıştırılması."),
        ("Pandas / OpenPyXL", "Excel girdilerinin okunması ve analiz sonuçlarının Excel dosyasına yazılması."),
        ("SQLite", "Analiz işlerinin durum, log, sonuç ve rapor bilgilerinin saklanması."),
        ("Chart.js", "Frontend üzerinde tema frekans grafiğinin gösterilmesi."),
        ("python-docx", "Akademik raporun Word biçiminde dışa aktarılması."),
    ], [4.0, 10.5])

    add_heading(doc, "Boru Hattı Aşamaları", 2)
    if len(diagrams) > 1:
        doc.add_picture(str(diagrams[1]), width=Inches(6.0))
        add_caption(doc, "Şekil 2. Beş Aşamalı Analiz Akışı")
    add_caption(doc, "Tablo 2. Uygulanan Analiz Boru Hattı Aşamaları")
    add_table(doc, ["Aşama", "Girdi", "İşlem", "Çıktı"], [
        ("1. Ingestion", "Excel/TXT/PDF dosyası", "Excel için text_data sütunu okunur; TXT/PDF için metin çıkarımı ve sliding window chunking uygulanır.", "Metin segmentleri"),
        ("2. Privacy Guard", "Ham veya segmentlenmiş metin", "Ollama qwen3:4b ile PII redaksiyonu denenir; hata durumunda regex PatternGuard çalışır.", "Maskelenmiş metin"),
        ("3. Theme Mapping", "Maskelenmiş metin", "Groq/Gemini tabanlı model havuzu üzerinden klinik tema ve semptom kodları çıkarılır.", "Tema listesi"),
        ("4. Validation", "Tema listesi ve metin", "PubMed ESearch/ESummary sorgusu ve LLM tabanlı denetim yapılır.", "YES/NO/PARTIAL doğrulama"),
        ("5. Reduction", "Tüm satır sonuçları", "AcademicReducer doğrulanmış temaları akademik raporda sentezler.", "Markdown rapor ve indirilebilir çıktılar"),
    ], [3.0, 3.2, 5.6, 2.7])

    add_heading(doc, "Model Havuzu ve Kota Duyarlı Yönlendirme", 2)
    add_body_paragraph(doc, "Kod tabanında get_model_pool ve call_model_pool fonksiyonları ile tema çıkarma, validasyon ve sentez görevleri için ayrı model havuzları tanımlanmıştır. Havuzlar THEME_MODEL_POOL, VALIDATION_MODEL_POOL ve REDUCTION_MODEL_POOL çevre değişkenleriyle değiştirilebilir. Bir modelde kota, rate-limit, model bulunamama, timeout veya servis erişim hatası oluştuğunda sistem sıradaki modele geçer.")
    add_body_paragraph(doc, "Router olayları çalışma zamanı iş nesnesinde tutulur ve frontend tarafındaki Router & Gizlilik sekmesinde kullanıcıya gösterilir. Bu panel, hangi modelin denendiğini, hangi modelde hata alındığını ve hangi modelin başarılı olduğunu izlenebilir hale getirir.")

    add_heading(doc, "Yerel Ollama Katmanı", 2)
    add_body_paragraph(doc, "Yerel modelin görevi ağır tıbbi yorumlama veya akademik sentez değildir. AgentA_PrivacyScrubber sınıfı, Ollama qwen3:4b modelini kişisel veri redaksiyonu için kullanır. Ayrıca uygulama başlangıcında değil, analiz işi başlarken ensure_ollama_service fonksiyonu ile Ollama API erişimi kontrol edilir; kapalıysa ollama serve komutu arka planda başlatılmaya çalışılır.")
    add_body_paragraph(doc, "Ollama'nın çalışmaması durumunda sistem hataya düşmek yerine TC kimlik numarası, telefon ve e-posta gibi desenleri yakalayan PatternGuard katmanına döner. Bu nedenle sistemde gizlilik katmanı iki parçalıdır: yerel LLM tabanlı redaksiyon denemesi ve deterministik regex tabanlı emniyet ağı.")

    add_heading(doc, "API Tasarımı", 2)
    add_caption(doc, "Tablo 3. API Uç Noktaları ve İşlevleri")
    add_table(doc, ["Uç Nokta", "Yöntem", "İşlev"], [
        ("/upload", "POST", "Dosyayı yükler, job_id oluşturur ve arka plan analiz işini başlatır."),
        ("/jobs", "GET", "SQLite veritabanında kayıtlı analiz işlerini listeler."),
        ("/status/{job_id}", "GET", "İş durumu, ilerleme, loglar, router olayları ve metrikleri döndürür."),
        ("/results/{job_id}", "GET", "Tamamlanan analizin satır sonuçlarını ve raporunu JSON olarak döndürür."),
        ("/download/{job_id}", "GET", "Analiz sonucunu Excel dosyası olarak indirir."),
        ("/download/report/{job_id}", "GET", "Akademik raporu Markdown dosyası olarak indirir."),
        ("/download/report/docx/{job_id}", "GET", "Akademik raporu Word dosyası olarak üretir ve indirir."),
        ("/cancel/{job_id}", "POST", "Kuyrukta, çalışan veya reduce aşamasındaki işi iptal eder."),
        ("/health", "GET", "Backend servisinin çalıştığını bildirir."),
    ], [4.3, 2.0, 8.2])

    add_heading(doc, "Üçüncü Bölüm", 1)
    add_heading(doc, "Uygulama, Deneysel Çalışma ve Değerlendirme", 2)
    add_body_paragraph(doc, "Projenin güncel uygulama davranışı app.py dosyasındaki FastAPI sürümü üzerinden gözlenmektedir. main.py dosyası aynı fikrin komut satırı prototipi olarak tutulmuştur. Web sürümü, kullanıcıya dosya yükleme ve canlı analiz takibi sağladığı için nihai prototip davranışı açısından daha temsil edicidir.")
    add_body_paragraph(doc, "Analiz işi FastAPI BackgroundTasks ile arka planda başlatılır. run_pipeline_job fonksiyonu işin durumunu running, reducing, done, error veya cancelled olarak günceller. Segment işleme aşamasında asyncio.Semaphore(2) kullanılarak aynı anda en fazla iki segmentin işlenmesi sağlanır; her segment öncesinde API limitlerini yumuşatmak için kısa bekleme uygulanır.")

    add_heading(doc, "Demo-Safe Çalışma Modu", 2)
    add_body_paragraph(doc, "Ücretsiz API kullanımında kota dolması veya model erişim hatası pratik bir risktir. Bu nedenle tema çıkarma, validasyon ve rapor sentezi aşamalarında demo-safe fallback mekanizmaları eklenmiştir. fallback_extract_themes temel anahtar kelime kurallarıyla tema üretir; fallback_validate_themes PubMed sonucu ve tema varlığına göre kısmi doğrulama döndürür; fallback_synthesize_report ise mevcut ara çıktılardan temel Markdown raporu oluşturur.")
    add_body_paragraph(doc, "Bu mekanizma akademik doğruluk iddiası taşımaz; sistemin demo ve savunma anında tamamen durmasını engelleyen dayanıklılık katmanı olarak değerlendirilmelidir.")

    add_heading(doc, "Arayüz ve Gözlemlenebilirlik", 2)
    add_body_paragraph(doc, "Frontend tarafında kullanıcı geçmiş analizleri görebilir, yeni Excel dosyası yükleyebilir, analiz sırasında terminal benzeri canlı log panelini izleyebilir ve işlem tamamlandığında veri tablosu, grafik, router-gizlilik paneli ve akademik rapor sekmeleri arasında geçiş yapabilir. Arayüzde sonuçlar Excel, Markdown ve Word dosyası olarak indirilebilir.")
    add_body_paragraph(doc, "Router & Gizlilik sekmesi; toplam PII redaksiyonu, PII bulunan segment sayısı, fallback kullanımı ve rapor modelini gösterir. Satır bazlı sonuç tablosunda tema çıkarma ve validasyon için kullanılan model etiketleri de gösterilir.")

    add_heading(doc, "Sınırlılıklar", 2)
    add_caption(doc, "Tablo 4. Sınırlılıklar ve Risk Azaltma Yaklaşımları")
    add_table(doc, ["Sınırlılık", "Etkisi", "Risk Azaltma"], [
        ("Klinik doğrulama uzman değerlendirmesiyle yapılmamıştır.", "Çıktılar klinik karar için doğrudan kullanılamaz.", "Rapor prototip olarak konumlandırılmış; uzman doğrulaması gelecek çalışma olarak önerilmiştir."),
        ("Regex PatternGuard yalnızca TC, telefon ve e-posta desenlerini kesin yakalar.", "İsim/adres gibi bağlamsal PII kaçabilir.", "Ollama redaksiyon katmanı ve çıktıdaki PII redaksiyon metrikleri kullanılmıştır."),
        ("Frontend dosya seçimi Excel odaklıdır.", "TXT/PDF desteği backend seviyesinde kalsa da arayüzde öne çıkarılmamıştır.", "Arayüz metinleri Excel iş akışına göre düzenlenmiştir."),
        ("Router olayları çalışma zamanı nesnesinde tutulur.", "Sunucu yeniden başladığında yeni router metrikleri kalıcı olmayabilir.", "Ana job durumu, log, sonuç ve rapor SQLite içinde saklanmaktadır."),
        ("Ücretsiz API kotası değişkenlik gösterir.", "Bazı analizlerde Gemini veya Groq kota hatası oluşabilir.", "Model havuzu ve demo-safe fallback uygulanmıştır."),
        ("PubMed sorguları ağ/sertifika sorunlarından etkilenebilir.", "Doğrulama aşaması PubMed başlığı bulamayabilir.", "No PubMed articles found durumunda LLM doğrulama ve fallback davranışı kullanılır."),
    ], [4.0, 4.7, 5.8])

    add_heading(doc, "Değerlendirme", 2)
    add_body_paragraph(doc, "Kod incelemesi, sistemin hedeflenen beş aşamalı analiz akışını uçtan uca yürütebildiğini göstermektedir. Dosya yükleme, iş takibi, segment işleme, model havuzu ile tema çıkarma, PubMed destekli doğrulama, akademik sentez, çıktı indirme ve iş iptali davranışları uygulama kodunda mevcuttur.")
    add_body_paragraph(doc, "Bununla birlikte, raporda F1-score, precision veya recall gibi nicel başarı metrikleri ölçülmüş gibi sunulmamıştır. Proje dokümanlarında hedef olarak belirtilen metrikler bu raporda gelecek çalışma ve değerlendirme hedefi olarak ele alınmıştır.")

    add_heading(doc, "Sonuç ve Öneriler", 1)
    add_body_paragraph(doc, "Bu çalışmada, tıp dikeyindeki nitel metin analizleri için Med-AgentLab adlı çoklu ajan orkestrasyon prototipi geliştirilmiş ve yazılım mimarisi ortaya konmuştur. Sistem, yerel gizlilik katmanı, model havuzu, PubMed destekli doğrulama ve rapor üretimi gibi modülleri bir araya getirerek tekil model kullanımının doğurduğu bazı riskleri azaltmayı amaçlamaktadır.")
    add_body_paragraph(doc, "Projenin en önemli mühendislik katkısı, görevlerin modele göre değil işlevsel aşamalara göre ayrılmasıdır. Ollama dar kapsamlı yerel gizlilik görevinde kullanılırken, Groq ve Gemini gibi bulut sağlayıcıları daha ağır tema çıkarma ve validasyon görevlerinde model havuzu içinde değerlendirilmiştir. Kota veya erişim hatalarında fallback yaklaşımı sistemin çalışmaya devam etmesini sağlar.")
    add_body_paragraph(doc, "Gelecek çalışmalarda gerçek uzman kodlamasıyla nicel doğruluk ölçümü yapılması, PII redaksiyon başarısının ayrı bir test veri setiyle değerlendirilmesi, router olaylarının kalıcı veritabanı şemasına eklenmesi, TXT/PDF yükleme desteğinin arayüzde görünür hale getirilmesi ve klinik güvenlik açısından daha sıkı denetim mekanizmalarının eklenmesi önerilmektedir.")


def add_references_and_appendices(doc: Document):
    add_heading(doc, "Kaynakça", 1)
    refs = [
        "Braun, V. ve Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77-101.",
        "Dean, J. ve Ghemawat, S. (2004). MapReduce: Simplified data processing on large clusters. OSDI'04: Sixth Symposium on Operating System Design and Implementation, 137-150.",
        "FastAPI. (2026). Background Tasks ve dosya yükleme dokümantasyonu. https://fastapi.tiangolo.com/",
        "LiteLLM. (2026). LiteLLM documentation: unified LLM access, retry and fallback logic. https://docs.litellm.ai/",
        "NCBI. (2026). Entrez Programming Utilities Help: ESearch and ESummary. https://www.ncbi.nlm.nih.gov/books/NBK25499/",
        "Ollama. (2026). Ollama API documentation. https://docs.ollama.com/api/introduction",
        "Samsun Üniversitesi. (2023). Bitirme Projesi Yazım Kılavuzu ve Şablonu.",
        "Tenacity. (2026). Tenacity retrying library documentation. https://tenacity.readthedocs.io/",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    break_page(doc)

    add_heading(doc, "Ekler", 1)
    add_heading(doc, "Ek 1. İncelenen Proje Dosyaları", 2)
    add_bullets(doc, [
        "app.py: FastAPI backend, model router, ajan sınıfları, fallback mekanizmaları ve çıktı uç noktaları.",
        "main.py: Komut satırı tabanlı ilk pipeline prototipi.",
        "frontend/index.html: Tek sayfalık web arayüzü, canlı log, sonuç tablosu, grafik ve router paneli.",
        "README.md, ABOUT.md, idea.md, Knowledge-Base/LLMs.txt: Proje vizyonu, mimari hedefler ve bilgi tabanı.",
        "requirements.txt: Kullanılan Python bağımlılıkları.",
    ])
    add_heading(doc, "Ek 2. Kurulum Notları", 2)
    add_bullets(doc, [
        "Python bağımlılıkları requirements.txt dosyası ile kurulmaktadır.",
        "Yerel model için Ollama kurulmalı ve qwen3:4b modeli indirilmelidir.",
        "Groq ve Gemini anahtarları .env dosyası üzerinden sağlanmalıdır.",
        "Web uygulaması uvicorn app:app --host 127.0.0.1 --port 8000 komutuyla çalıştırılabilir.",
    ])
    break_page(doc)

    add_heading(doc, "Özgeçmiş", 1)
    add_body_paragraph(doc, f"{STUDENT}, Samsun Üniversitesi {DEPARTMENT} öğrencisidir. Yazılım mühendisliği eğitimi kapsamında yapay zeka, web tabanlı sistem geliştirme, API entegrasyonları ve çoklu model orkestrasyonu alanlarında çalışmalar yürütmektedir. Bu bitirme projesinde, kısıtlı donanım koşullarında yerel ve bulut tabanlı yapay zeka modellerini birlikte kullanabilen bir nitel veri analizi prototipi geliştirmiştir.")
    add_body_paragraph(doc, "Yabancı Dil: [Doldurulacak]")
    add_body_paragraph(doc, "E-Posta: [Doldurulacak]")


def build():
    diagrams = create_diagrams()
    doc = Document()
    apply_document_styles(doc)

    add_cover(doc)
    break_page(doc)
    add_inner_cover(doc)
    break_page(doc)
    add_approval(doc)
    break_page(doc)
    add_declaration(doc)

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    front.left_margin = Cm(3)
    front.right_margin = Cm(2.5)
    front.top_margin = Cm(2.5)
    front.bottom_margin = Cm(2.5)
    add_page_number(front, fmt="upperRoman", start=4)
    add_front_matter(doc)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.left_margin = Cm(3)
    body.right_margin = Cm(2.5)
    body.top_margin = Cm(2.5)
    body.bottom_margin = Cm(2.5)
    add_page_number(body, fmt="decimal", start=1)
    add_main_text(doc, diagrams)
    add_references_and_appendices(doc)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Bitirme Projesi Final Raporu"
    doc.core_properties.author = STUDENT
    doc.core_properties.keywords = "Med-AgentLab, LLM, MapReduce, Nitel Veri Analizi, Gizlilik"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
