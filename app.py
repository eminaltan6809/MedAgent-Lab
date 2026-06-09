#!/usr/bin/env python3
"""
FastAPI Backend for Zero-Cost Multi-Agent Qualitative Data Analysis Pipeline (Med-AgentLab)
Supports 5 stages w/ background processing, Excel/TXT/PDF upload, and synthesized report endpoints.
"""

import asyncio
import logging
import os
import uuid
import sys
import re
import subprocess
import shutil
import time
import urllib.request
import urllib.parse
import json
from dataclasses import dataclass
from typing import List, Dict, Any, Optional
from pathlib import Path

import pandas as pd
from fastapi import FastAPI, File, UploadFile, BackgroundTasks, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
from litellm import acompletion
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# Load Environment
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('analysis_pipeline.log')
    ]
)
logger = logging.getLogger(__name__)

app = FastAPI(title="Med-AgentLab Qualitative Analysis API")

# Enable CORS for frontend requests
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory job store
jobs: Dict[str, Dict[str, Any]] = {}

# ---------------------------------------------------------------------------
# SQLite Database Integration
# ---------------------------------------------------------------------------
import sqlite3

DB_PATH = "med_agentlab.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS jobs (
            job_id TEXT PRIMARY KEY,
            filename TEXT,
            status TEXT,
            progress INTEGER,
            processed INTEGER,
            total INTEGER,
            error TEXT,
            report TEXT,
            results TEXT,
            logs TEXT,
            created_at TEXT
        )
    """)
    cursor.execute("PRAGMA table_info(jobs)")
    existing_columns = {row[1] for row in cursor.fetchall()}
    meta_columns = {
        "router_events": "TEXT",
        "privacy_metrics": "TEXT",
        "model_summary": "TEXT",
        "report_model": "TEXT",
        "report_fallback_used": "INTEGER DEFAULT 0",
    }
    for column, column_type in meta_columns.items():
        if column not in existing_columns:
            cursor.execute(f"ALTER TABLE jobs ADD COLUMN {column} {column_type}")
    conn.commit()
    conn.close()

def summarize_job_metadata(results: list, total: int = 0) -> dict:
    total_redactions = sum(int(r.get("pii_redaction_count", 0) or 0) for r in results)
    theme_fallback_count = sum(1 for r in results if r.get("theme_fallback_used"))
    validation_fallback_count = sum(1 for r in results if r.get("validation_fallback_used"))
    return {
        "privacy_metrics": {
            "total_redactions": total_redactions,
            "segments_with_redactions": sum(1 for r in results if int(r.get("pii_redaction_count", 0) or 0) > 0),
            "total_segments": total or len(results),
        },
        "model_summary": {
            "theme_models": sorted({r.get("theme_model") for r in results if r.get("theme_model")}),
            "validation_models": sorted({r.get("validation_model") for r in results if r.get("validation_model")}),
            "theme_fallback_count": theme_fallback_count,
            "validation_fallback_count": validation_fallback_count,
            "fallback_count": theme_fallback_count + validation_fallback_count,
        },
    }

def reconstruct_router_events(logs: list) -> list:
    events = []
    for entry in logs or []:
        message = str(entry.get("message", ""))
        attempt = re.search(r"Model havuzu denemesi \(([^)]+)\) \d+/\d+: (.+)$", message)
        success = re.search(r"Model havuzu basarili \(([^)]+)\): (.+)$", message)
        if attempt:
            events.append({
                "time": entry.get("time", ""),
                "task": attempt.group(1),
                "model": attempt.group(2),
                "status": "attempt",
                "message": message,
            })
        elif success:
            events.append({
                "time": entry.get("time", ""),
                "task": success.group(1),
                "model": success.group(2),
                "status": "success",
                "message": message,
            })
        elif "API havuzu kullanilamadi" in message or "demo-safe" in message.lower():
            task = "reduction" if entry.get("agent") == "agent_d" else "unknown"
            events.append({
                "time": entry.get("time", ""),
                "task": task,
                "model": "demo-safe/local-fallback",
                "status": "fallback_success",
                "message": message,
            })
    return events

def save_job_to_db(job_id: str, job_data: dict):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO jobs (
            job_id, filename, status, progress, processed, total, error, report,
            results, logs, created_at, router_events, privacy_metrics,
            model_summary, report_model, report_fallback_used
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(job_id) DO UPDATE SET
            filename=excluded.filename,
            status=excluded.status,
            progress=excluded.progress,
            processed=excluded.processed,
            total=excluded.total,
            error=excluded.error,
            report=excluded.report,
            results=excluded.results,
            logs=excluded.logs,
            router_events=excluded.router_events,
            privacy_metrics=excluded.privacy_metrics,
            model_summary=excluded.model_summary,
            report_model=excluded.report_model,
            report_fallback_used=excluded.report_fallback_used
    """, (
        job_id,
        job_data.get("filename", ""),
        job_data.get("status", ""),
        job_data.get("progress", 0),
        job_data.get("processed", 0),
        job_data.get("total", 0),
        job_data.get("error"),
        job_data.get("report", ""),
        json.dumps(job_data.get("results", [])),
        json.dumps(job_data.get("logs", [])),
        job_data.get("created_at", ""),
        json.dumps(job_data.get("router_events", [])),
        json.dumps(job_data.get("privacy_metrics", {})),
        json.dumps(job_data.get("model_summary", {})),
        job_data.get("report_model", ""),
        1 if job_data.get("report_fallback_used", False) else 0,
    ))
    conn.commit()
    conn.close()

def save_job(job_id: str):
    if job_id in jobs:
        save_job_to_db(job_id, jobs[job_id])

def load_jobs_from_db():
    try:
        init_db()
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT job_id, filename, status, progress, processed, total, error,
                   report, results, logs, created_at, router_events,
                   privacy_metrics, model_summary, report_model,
                   report_fallback_used
            FROM jobs
        """)
        rows = cursor.fetchall()
        conn.close()
        for row in rows:
            results = json.loads(row[8]) if row[8] else []
            logs = json.loads(row[9]) if row[9] else []
            derived = summarize_job_metadata(results, row[5] or 0)
            router_events = json.loads(row[11]) if len(row) > 11 and row[11] else reconstruct_router_events(logs)
            privacy_metrics = json.loads(row[12]) if len(row) > 12 and row[12] else derived["privacy_metrics"]
            model_summary = json.loads(row[13]) if len(row) > 13 and row[13] else derived["model_summary"]
            jobs[row[0]] = {
                "status": row[2],
                "progress": row[3],
                "processed": row[4],
                "total": row[5],
                "error": row[6],
                "report": row[7],
                "results": results,
                "logs": logs,
                "filename": row[1] if row[1] else "",
                "created_at": row[10] if row[10] else "",
                "output_path": f"outputs/{row[0]}_output.xlsx",
                "router_events": router_events,
                "privacy_metrics": privacy_metrics,
                "model_summary": model_summary,
                "report_model": row[14] if len(row) > 14 and row[14] else "",
                "report_fallback_used": bool(row[15]) if len(row) > 15 else False,
            }
        logger.info(f"Veritabani yuklendi. {len(rows)} eski analiz yuklendi.")
    except Exception as e:
        logger.error(f"Veritabani yukleme hatasi: {e}")

# Initialize and Load
load_jobs_from_db()

# ---------------------------------------------------------------------------
# Logging helper for client communication
# ---------------------------------------------------------------------------
from datetime import datetime

def log_job(job_id: Optional[str], message: str, agent: str = "system", is_error: bool = False):
    if not job_id:
        return
    if job_id in jobs:
        if "logs" not in jobs[job_id]:
            jobs[job_id]["logs"] = []
        timestamp = datetime.now().strftime("%H:%M:%S")
        entry = {
            "time": timestamp,
            "agent": agent,
            "message": message,
            "is_error": is_error
        }
        jobs[job_id]["logs"].append(entry)
        tag = f"[{agent.upper()}]"
        if is_error:
            logger.error(f"{tag} Job {job_id}: {message}")
        else:
            logger.info(f"{tag} Job {job_id}: {message}")
        save_job(job_id)

def record_router_event(
    job_id: Optional[str],
    task_name: str,
    model: str,
    status: str,
    message: str = "",
):
    if not job_id or job_id not in jobs:
        return
    jobs[job_id].setdefault("router_events", [])
    jobs[job_id]["router_events"].append({
        "time": datetime.now().strftime("%H:%M:%S"),
        "task": task_name,
        "model": model,
        "status": status,
        "message": message,
    })
    save_job(job_id)

# ---------------------------------------------------------------------------
# Ollama service bootstrap
# ---------------------------------------------------------------------------

def get_ollama_base_url() -> str:
    return os.getenv("OLLAMA_API_BASE", "http://127.0.0.1:11434").rstrip("/")

def get_ollama_model_name() -> str:
    model = os.getenv("OLLAMA_MODEL", "ollama/qwen3:4b")
    return model.replace("ollama/", "", 1)

def is_ollama_api_ready() -> bool:
    try:
        with urllib.request.urlopen(f"{get_ollama_base_url()}/api/tags", timeout=2) as response:
            return response.status == 200
    except Exception:
        return False

def list_ollama_models() -> List[str]:
    try:
        with urllib.request.urlopen(f"{get_ollama_base_url()}/api/tags", timeout=3) as response:
            data = json.loads(response.read().decode())
            return [m.get("name", "") for m in data.get("models", []) if m.get("name")]
    except Exception:
        return []

def ensure_ollama_service_sync(job_id: Optional[str] = None) -> bool:
    model_name = get_ollama_model_name()
    if is_ollama_api_ready():
        models = list_ollama_models()
        if model_name not in models:
            log_job(job_id, f"Ollama calisiyor fakat '{model_name}' modeli bulunamadi. 'ollama pull {model_name}' gerekebilir.", "agent_a", is_error=True)
        else:
            log_job(job_id, f"Ollama hazir. Yerel privacy modeli bulundu: {model_name}", "agent_a")
        return True

    ollama_exe = shutil.which("ollama")
    if not ollama_exe:
        log_job(job_id, "Ollama komutu bulunamadi. Regex Pattern Guard ile devam edilecek.", "agent_a", is_error=True)
        return False

    log_job(job_id, "Ollama kapali gorunuyor. Analiz oncesi 'ollama serve' baslatiliyor...", "agent_a")
    try:
        creationflags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
        subprocess.Popen(
            [ollama_exe, "serve"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            creationflags=creationflags,
        )
    except Exception as e:
        log_job(job_id, f"Ollama otomatik baslatilamadi: {str(e)}. Regex Pattern Guard ile devam edilecek.", "agent_a", is_error=True)
        return False

    for _ in range(15):
        time.sleep(1)
        if is_ollama_api_ready():
            models = list_ollama_models()
            if model_name not in models:
                log_job(job_id, f"Ollama baslatildi ancak '{model_name}' modeli yok. 'ollama pull {model_name}' gerekebilir.", "agent_a", is_error=True)
            else:
                log_job(job_id, f"Ollama baslatildi ve hazir: {model_name}", "agent_a")
            return True

    log_job(job_id, "Ollama baslatma denemesi zaman asimina ugradi. Regex Pattern Guard ile devam edilecek.", "agent_a", is_error=True)
    return False

async def ensure_ollama_service(job_id: Optional[str] = None) -> bool:
    return await asyncio.to_thread(ensure_ollama_service_sync, job_id)

# ---------------------------------------------------------------------------
# Quota-aware model routing
# ---------------------------------------------------------------------------

def _csv_env(name: str, defaults: List[str]) -> List[str]:
    raw = os.getenv(name, "")
    values = [item.strip() for item in raw.split(",") if item.strip()]
    pool = values or defaults
    deduped = []
    for model in pool:
        if model not in deduped:
            deduped.append(model)
    return deduped

def get_model_pool(task_name: str) -> List[str]:
    groq_model = os.getenv("GROQ_MODEL", "groq/llama-3.1-8b-instant")
    gemini_fast = os.getenv("GEMINI_FAST_MODEL", "gemini/gemini-1.5-flash")
    gemini_lite = os.getenv("GEMINI_LITE_MODEL", "gemini/gemini-2.0-flash-lite")

    pools = {
        "theme_mapping": _csv_env(
            "THEME_MODEL_POOL",
            [groq_model, gemini_fast, gemini_lite],
        ),
        "validation": _csv_env(
            "VALIDATION_MODEL_POOL",
            [os.getenv("AGENT_C_MODEL", gemini_lite), gemini_fast, groq_model],
        ),
        "reduction": _csv_env(
            "REDUCTION_MODEL_POOL",
            [os.getenv("AGENT_D_MODEL", gemini_lite), gemini_fast, groq_model],
        ),
    }
    return pools[task_name]

def is_pool_switch_error(exc: Exception) -> bool:
    text = f"{type(exc).__name__}: {exc}".lower()
    switch_markers = [
        "ratelimit",
        "rate limit",
        "quota",
        "resource_exhausted",
        "429",
        "notfound",
        "model not found",
        "model_not_found",
        "badrequest",
        "invalid model",
        "timeout",
        "temporarily unavailable",
        "service unavailable",
    ]
    return any(marker in text for marker in switch_markers)

async def call_model_pool(
    task_name: str,
    messages: List[Dict[str, str]],
    job_id: Optional[str] = None,
    agent: str = "system",
    **kwargs,
):
    last_error: Optional[Exception] = None
    models = get_model_pool(task_name)

    for idx, model in enumerate(models, start=1):
        try:
            log_job(job_id, f"Model havuzu denemesi ({task_name}) {idx}/{len(models)}: {model}", agent)
            record_router_event(job_id, task_name, model, "attempt", f"Deneme {idx}/{len(models)}")
            response = await acompletion(model=model, messages=messages, **kwargs)
            log_job(job_id, f"Model havuzu basarili ({task_name}): {model}", agent)
            status = "primary_success" if idx == 1 else "fallback_success"
            record_router_event(job_id, task_name, model, status, "Model yanit verdi")
            return response, model
        except Exception as e:
            last_error = e
            if idx < len(models) and is_pool_switch_error(e):
                record_router_event(job_id, task_name, model, "failed_over", str(e)[:220])
                log_job(
                    job_id,
                    f"{model} kullanilamadi/kotaya takildi. Siradaki modele geciliyor: {str(e)[:220]}",
                    agent,
                    is_error=True,
                )
                continue
            record_router_event(job_id, task_name, model, "failed", str(e)[:220])
            raise

    if last_error:
        raise last_error
    raise RuntimeError(f"Model pool is empty for task: {task_name}")

# ---------------------------------------------------------------------------
# Data classes and helper functions
# ---------------------------------------------------------------------------

@dataclass
class AnalysisResult:
    original_text: str
    pii_redacted_text: str
    mapped_themes: str
    validation_result: str
    validation_reason: str
    theme_model: str = ""
    validation_model: str = ""
    theme_fallback_used: bool = False
    validation_fallback_used: bool = False
    pii_redaction_count: int = 0

def chunk_text_sliding_window(text: str, max_words: int = 2250, overlap_words: int = 340) -> List[str]:
    """Splits text into chunks of roughly 3000 tokens with 15% overlap"""
    words = text.split()
    if len(words) <= max_words:
        return [text]
    
    chunks = []
    start = 0
    while start < len(words):
        end = start + max_words
        chunk_words = words[start:end]
        chunks.append(" ".join(chunk_words))
        if end >= len(words):
            break
        start += (max_words - overlap_words)
    return chunks

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extracts text from PDF file using PyPDF2 if installed"""
    try:
        import PyPDF2
        text = ""
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except ImportError:
        raise HTTPException(
            status_code=400, 
            detail="Sistemde PDF işlemek için PyPDF2 kütüphanesi yüklü değil. Lütfen 'pip install PyPDF2' komutunu çalıştırın veya TXT/Excel yükleyin."
        )

def count_redaction_tags(text: str) -> int:
    return len(re.findall(r"\[[A-Z_]+_REDACTED\]", text or ""))

def fallback_extract_themes(text: str) -> str:
    lower = (text or "").lower()
    theme_rules = [
        ("depresyon", ["depresyon", "hads-d", "beck", "ağlama", "ölüm düşünceleri", "intihar"]),
        ("anksiyete", ["anksiyete", "panik", "kalp çarp", "nefes alam", "kabus"]),
        ("uyku bozukluğu", ["uyku", "uyuyam", "insomnia", "kabus"]),
        ("kronik ağrı", ["ağrı", "fibromiyalji", "bel ağrısı", "tender"]),
        ("yorgunluk", ["yorgun", "halsiz", "bitkin"]),
        ("sosyal izolasyon", ["sosyal izolasyon", "dışarı", "avm", "toplu taşıma", "arkadaş"]),
        ("travma belirtileri", ["travma", "flashback", "kaza", "irkilme", "kaçınma"]),
        ("bağımlılık", ["alkol", "benzodiazepin", "detoks", "yoksunluk", "audit"]),
        ("bilişsel bozulma", ["unutkan", "demans", "mmse", "dezoryantasyon", "alzheimer"]),
        ("tedavi uyumu", ["ilaç uyumsuz", "tedavi", "terapi", "takip", "program"]),
        ("işlevsellik kaybı", ["iş", "okul", "performans", "günlük aktivite", "yaşam kalitesi"]),
        ("beslenme sorunu", ["iştah", "kilo", "bmi", "anoreksiya", "yemek"]),
    ]
    themes = [theme for theme, needles in theme_rules if any(n in lower for n in needles)]
    return ", ".join(themes[:8] or ["klinik stres", "yaşam kalitesi etkilenimi", "tedavi ihtiyacı"])

def fallback_validate_themes(themes: str, pubmed_titles: List[str]) -> tuple[str, str]:
    if pubmed_titles:
        return "PARTIAL", "Demo-safe doğrulama kullanıldı; PubMed başlıkları bulunduğu için temalar kısmen destekli işaretlendi."
    if themes and themes != "ERROR":
        return "PARTIAL", "Demo-safe doğrulama kullanıldı; API havuzu erişilemediği için klinik açıdan makul temalar kısmi olarak işaretlendi."
    return "ERROR", "Doğrulanacak tema üretilemedi."

def fallback_synthesize_report(results: List[AnalysisResult]) -> str:
    counts: Dict[str, int] = {}
    valid_count = 0
    for res in results:
        if res.mapped_themes == "ERROR":
            continue
        if "YES" in res.validation_result.upper() or "PARTIAL" in res.validation_result.upper():
            valid_count += 1
        for theme in res.mapped_themes.split(","):
            clean = theme.strip()
            if clean:
                counts[clean] = counts.get(clean, 0) + 1

    top_themes = sorted(counts.items(), key=lambda x: x[1], reverse=True)
    theme_lines = "\n".join([f"- **{theme}**: {count} segment" for theme, count in top_themes[:12]]) or "- Tema üretilemedi."
    return (
        "# Med-AgentLab Demo-Safe Akademik Rapor\n\n"
        "## Executive Summary\n"
        f"Bu rapor, API havuzu erişimi/kota sınırı nedeniyle deterministik demo-safe reducer tarafından üretildi. "
        f"Toplam {len(results)} segmentin {valid_count} tanesi doğrulanmış veya kısmi doğrulanmış tema içeriyor.\n\n"
        "## Codebook / Theme Tree\n"
        f"{theme_lines}\n\n"
        "## Discussion\n"
        "Bulgular klinik görüşmelerde psikolojik yük, işlevsellik kaybı, tedavi ihtiyacı ve yaşam kalitesi etkilenimi etrafında kümelenmektedir. "
        "Bu çıktı demo sürekliliği için oluşturulmuştur; akademik kullanımda uzman değerlendirmesi ve tam LLM senteziyle desteklenmelidir.\n\n"
        "## Conclusion\n"
        "Model havuzu kullanılamadığında sistem analizi tamamen durdurmak yerine yapılandırılmış ara çıktılardan okunabilir bir ön rapor üretmiştir."
    )

# ---------------------------------------------------------------------------
# Regex Pattern Guard for PII
# ---------------------------------------------------------------------------

class PatternGuard:
    def __init__(self):
        self.patterns = {
            "TC_NO": r"\b[1-9]\d{10}\b",
            "PHONE": r"\b(?:\+?90[- ]?)?0?[5-9]\d{2}[- ]?\d{3}[- ]?\d{4}\b",
            "EMAIL": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        }
        
    def sanitize(self, text: str) -> str:
        sanitized = text
        for label, pattern in self.patterns.items():
            sanitized = re.sub(pattern, f"[{label}_REDACTED]", sanitized)
        return sanitized

# ---------------------------------------------------------------------------
# Multi-Agent Classes
# ---------------------------------------------------------------------------

class AgentA_PrivacyScrubber:
    def __init__(self):
        self.model = os.getenv("OLLAMA_MODEL", "ollama/qwen3:4b")
        self.pattern_guard = PatternGuard()
        self.system_prompt = (
            "You are a lightweight local privacy and preprocessing guard. "
            "Read the clinical text/interview and redact obvious Personally Identifiable Information (PII) "
            "such as names, surnames, ID numbers, phone numbers, emails, and addresses. "
            "Replace them with appropriate tags like [NAME_REDACTED] or [ADDRESS_REDACTED]. "
            "Do not analyze, summarize, classify, or rewrite clinical meaning. "
            "Do NOT change medical details, symptoms, scales, medications, or other clinical content. "
            "Return ONLY the redacted text without any introductory or concluding comments."
        )

    @retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=2, max=10),
           retry=retry_if_exception_type(Exception))
    async def scrub_pii(self, text: str, row_num: int, job_id: Optional[str] = None) -> str:
        log_job(job_id, f"Ajan A (PrivacyScrubber) calismaya basladi. Satir {row_num}. Metin: '{text[:120]}...'", "agent_a")
        log_job(job_id, f"Ajan A Istemi (Prompt):\n[System]\n{self.system_prompt}\n[User]\nText to redact:\n{text[:200]}...", "agent_a")
        try:
            response = await acompletion(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": f"Text to redact:\n{text}"}
                ],
                temperature=0.0,
                max_tokens=2048
            )
            redacted = response.choices[0].message.content.strip()
            log_job(job_id, f"Ajan A LLM Yaniti: '{redacted[:120]}...'", "agent_a")
        except Exception as e:
            log_job(job_id, f"Ajan A LLM hatasi aldi: {str(e)}. Regex Pattern Guard devreye sokuluyor.", "agent_a", is_error=True)
            redacted = text
            
        redacted = self.pattern_guard.sanitize(redacted)
        log_job(job_id, f"Ajan A tamamlandi (Satir {row_num}). Maskelenmis Metin: '{redacted[:120]}...'", "agent_a")
        return redacted


class AgentB_ThematicMapper:
    def __init__(self):
        self.system_prompt = (
            "You are a clinical qualitative data coder. "
            "Extract psychological, physical, and behavioral themes, symptoms or side effects from the text. "
            "Provide them as a comma-separated list of short codes or keywords. "
            "Do NOT include conversational text. Return ONLY the comma-separated codes/themes."
        )

    @retry(stop=stop_after_attempt(5), wait=wait_exponential(multiplier=1.5, min=5, max=60),
           retry=retry_if_exception_type(Exception))
    async def extract_themes(self, text: str, row_num: int, job_id: Optional[str] = None) -> tuple[str, str, bool]:
        log_job(job_id, f"Ajan B (ThematicMapper) calismaya basladi. Satir {row_num}. Metin: '{text[:120]}...'", "agent_b")
        log_job(job_id, f"Ajan B Istemi (Prompt):\n[System]\n{self.system_prompt}\n[User]\nRedacted text:\n{text[:200]}...", "agent_b")
        try:
            response, used_model = await call_model_pool(
                "theme_mapping",
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": f"Redacted text:\n{text}"}
                ],
                job_id=job_id,
                agent="agent_b",
                temperature=0.2,
                max_tokens=250,
            )
            themes = response.choices[0].message.content.strip()
            log_job(job_id, f"Ajan B LLM Yaniti ({used_model}) (Temalar): '{themes}'", "agent_b")
            return themes, used_model, used_model != get_model_pool("theme_mapping")[0]
        except Exception as e:
            themes = fallback_extract_themes(text)
            log_job(
                job_id,
                f"Ajan B API havuzu kullanilamadi, demo-safe tema cikarimi devrede: '{themes}'",
                "agent_b",
                is_error=True,
            )
            record_router_event(job_id, "theme_mapping", "demo-safe/local-keyword", "fallback_success", str(e)[:220])
            return themes, "demo-safe/local-keyword", True


class AgentC_PubMedValidator:
    def __init__(self):
        self.system_prompt = (
            "You are a senior medical reviewer and auditor. "
            "Verify whether the extracted medical themes/symptoms correspond to known medical literature. "
            "We searched PubMed database and found some reference article titles for these themes. "
            "Analyze the themes and reference titles. Decide if the themes are valid or potential hallucinations. "
            "Reply exactly in this format:\n"
            "VALIDATION: YES (or NO or PARTIAL)\n"
            "REASON: [Write one brief sentence explaining why, referencing the PubMed titles if supportive]"
        )

    def sync_search_pubmed(self, term: str) -> List[str]:
        try:
            clean_term = re.sub(r'[^\w\s]', ' ', term).strip()
            if not clean_term:
                return []
            query = urllib.parse.quote(clean_term)
            url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?db=pubmed&term={query}&retmode=json&retmax=2"
            
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=4) as response:
                if response.status == 200:
                    data = json.loads(response.read().decode())
                    id_list = data.get("esearchresult", {}).get("idlist", [])
                    if not id_list:
                        return []
                    
                    ids = ",".join(id_list)
                    summary_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?db=pubmed&id={ids}&retmode=json"
                    
                    req_sum = urllib.request.Request(summary_url, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req_sum, timeout=4) as sum_response:
                        if sum_response.status == 200:
                            sum_data = json.loads(sum_response.read().decode())
                            results = sum_data.get("result", {})
                            titles = []
                            for uid in id_list:
                                title = results.get(uid, {}).get("title", "")
                                if title:
                                    titles.append(title)
                            return titles
        except Exception as e:
            logger.warning(f"PubMed search error: {e}")
        return []

    async def search_pubmed(self, term: str) -> List[str]:
        return await asyncio.to_thread(self.sync_search_pubmed, term)

    @retry(stop=stop_after_attempt(5), wait=wait_exponential(multiplier=1.5, min=5, max=60),
           retry=retry_if_exception_type(Exception))
    async def validate_themes(self, themes: str, original_text: str, row_num: int, job_id: Optional[str] = None) -> tuple[str, str, str, bool]:
        log_job(job_id, f"Ajan C (PubMedValidator) calismaya basladi. PubMed arama terimi: '{themes}'", "agent_c")
        pubmed_titles = await self.search_pubmed(themes)
        rag_context = "\n".join([f"- {t}" for t in pubmed_titles]) if pubmed_titles else "No PubMed articles found."
        log_job(job_id, f"PubMed Arama Sonuclari (Satir {row_num}):\n{rag_context}", "agent_c")
        
        user_content = (
            f"Original Text snippet:\n{original_text}\n\n"
            f"Extracted Themes: {themes}\n\n"
            f"PubMed Search:\n{rag_context}"
        )
        
        log_job(job_id, f"Ajan C Istemi (Prompt):\n[System]\n{self.system_prompt}\n[User]\n{user_content[:300]}...", "agent_c")
        
        try:
            response, used_model = await call_model_pool(
                "validation",
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": user_content}
                ],
                job_id=job_id,
                agent="agent_c",
                temperature=0.1,
                max_tokens=200,
            )
            output = response.choices[0].message.content.strip()
            log_job(job_id, f"Ajan C LLM Yaniti ({used_model}):\n{output}", "agent_c")
            
            val_result = "YES"
            val_reason = output
            
            for line in output.split("\n"):
                if line.upper().startswith("VALIDATION:"):
                    val_result = line.split(":", 1)[1].strip()
                elif line.upper().startswith("REASON:"):
                    val_reason = line.split(":", 1)[1].strip()
            return val_result, val_reason, used_model, used_model != get_model_pool("validation")[0]
        except Exception as e:
            val_result, val_reason = fallback_validate_themes(themes, pubmed_titles)
            log_job(
                job_id,
                f"Ajan C API havuzu kullanilamadi, demo-safe dogrulama devrede: {val_result} - {val_reason}",
                "agent_c",
                is_error=True,
            )
            record_router_event(job_id, "validation", "demo-safe/pubmed-heuristic", "fallback_success", str(e)[:220])
            return val_result, val_reason, "demo-safe/pubmed-heuristic", True


class AgentD_AcademicReducer:
    def __init__(self):
        self.system_prompt = (
            "You are the Lead Academic Researcher in a medical university laboratory. "
            "Your task is to synthesize the extracted and validated qualitative codes/themes into a cohesive, "
            "structured Thematic Analysis Report suitable for a medical journal publication. "
            "Group the codes into hierarchical categories, deduplicate terms, and explain the key clinical findings. "
            "Format the report beautifully in Markdown with clear sections (e.g., Executive Summary, Codebook/Theme Tree, Discussion, Conclusion)."
        )

    @retry(stop=stop_after_attempt(5), wait=wait_exponential(multiplier=1.5, min=5, max=60),
           retry=retry_if_exception_type(Exception))
    async def synthesize_report(self, results: List[AnalysisResult], job_id: Optional[str] = None) -> tuple[str, str, bool]:
        log_job(job_id, f"Ajan D (AcademicReducer) calismaya basladi. Sentezlenen girdi boyutu: {len(results)} satir.", "agent_d")
        data_summary = []
        for idx, res in enumerate(results):
            if res.mapped_themes == "ERROR":
                continue
            data_summary.append(
                f"Snippet {idx+1}:\n"
                f"- Themes: {res.mapped_themes}\n"
                f"- Validation: {res.validation_result} ({res.validation_reason})"
            )
        
        input_data = "\n\n".join(data_summary)
        log_job(job_id, f"Ajan D Istemi (Prompt):\n[System]\n{self.system_prompt}\n[User]\nSentezlenecek Veri:\n{input_data[:400]}...", "agent_d")
        
        try:
            response, used_model = await call_model_pool(
                "reduction",
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": f"Qualitative analysis data:\n\n{input_data}"}
                ],
                job_id=job_id,
                agent="agent_d",
                temperature=0.3,
                max_tokens=2048,
            )
            report = response.choices[0].message.content.strip()
            log_job(job_id, f"Ajan D sentezi tamamladi ({used_model}). Rapor boyutu: {len(report)} karakter.", "agent_d")
            return report, used_model, used_model != get_model_pool("reduction")[0]
        except Exception as e:
            report = fallback_synthesize_report(results)
            log_job(job_id, "Ajan D API havuzu kullanilamadi, demo-safe reducer raporu uretildi.", "agent_d", is_error=True)
            record_router_event(job_id, "reduction", "demo-safe/local-reducer", "fallback_success", str(e)[:220])
            return report, "demo-safe/local-reducer", True

# ---------------------------------------------------------------------------
# Pipeline class for app job runs
# ---------------------------------------------------------------------------

class QualitativeAnalysisPipeline:
    def __init__(self):
        self.agent_a = AgentA_PrivacyScrubber()
        self.agent_b = AgentB_ThematicMapper()
        self.agent_c = AgentC_PubMedValidator()
        self.agent_d = AgentD_AcademicReducer()
        self.results: List[AnalysisResult] = []
        self.report: str = ""
        self.report_model: str = ""
        self.report_fallback_used: bool = False

    async def process_single_text(self, text: str, row_num: int, job_id: Optional[str] = None) -> AnalysisResult:
        if job_id and jobs.get(job_id, {}).get("status") == "cancelled":
            raise asyncio.CancelledError("İş iptal edildi")
        try:
            # Stage 2: Privacy Guard
            redacted_text = await self.agent_a.scrub_pii(text, row_num, job_id)
            if job_id and jobs.get(job_id, {}).get("status") == "cancelled":
                raise asyncio.CancelledError("İş iptal edildi")
                
            # Stage 3: Map Phase
            themes, theme_model, theme_fallback_used = await self.agent_b.extract_themes(redacted_text, row_num, job_id)
            if job_id and jobs.get(job_id, {}).get("status") == "cancelled":
                raise asyncio.CancelledError("İş iptal edildi")
                
            # Stage 4: Validation
            val_result, val_reason, validation_model, validation_fallback_used = await self.agent_c.validate_themes(themes, redacted_text, row_num, job_id)
            
            return AnalysisResult(
                text,
                redacted_text,
                themes,
                val_result,
                val_reason,
                theme_model=theme_model,
                validation_model=validation_model,
                theme_fallback_used=theme_fallback_used,
                validation_fallback_used=validation_fallback_used,
                pii_redaction_count=count_redaction_tags(redacted_text),
            )
        except asyncio.CancelledError as e:
            raise e
        except Exception as e:
            log_job(job_id, f"Kritik satir hatasi (Satir {row_num}): {str(e)}", "system", is_error=True)
            return AnalysisResult(text, "ERROR", "ERROR", "ERROR", str(e))

    def save_results(self, output_file: str) -> None:
        data = {
            'Original Text': [r.original_text for r in self.results],
            'PII Redacted Text': [r.pii_redacted_text for r in self.results],
            'Mapped Themes': [r.mapped_themes for r in self.results],
            'Validation Result': [r.validation_result for r in self.results],
            'Validation Reason': [r.validation_reason for r in self.results],
            'Theme Model': [r.theme_model for r in self.results],
            'Validation Model': [r.validation_model for r in self.results],
            'Theme Fallback Used': [r.theme_fallback_used for r in self.results],
            'Validation Fallback Used': [r.validation_fallback_used for r in self.results],
            'PII Redaction Count': [r.pii_redaction_count for r in self.results],
        }
        df = pd.DataFrame(data)
        df.to_excel(output_file, index=False, engine='openpyxl')
        
        # Save academic report
        report_file = output_file.replace("_output.xlsx", "_report.md").replace(".xlsx", "_report.md")
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write(self.report)

# ---------------------------------------------------------------------------
# Background job runner
# ---------------------------------------------------------------------------

async def run_pipeline_job(job_id: str, input_path: str, output_path: str, is_excel: bool):
    try:
        log_job(job_id, f"Analiz baslatiliyor... Girdi tipi: {'Excel' if is_excel else 'Transkript Metni'}", "system")
        jobs[job_id]["status"] = "running"
        save_job(job_id)
        await ensure_ollama_service(job_id)

        if is_excel:
            df = pd.read_excel(input_path)
            if 'text_data' not in df.columns:
                raise ValueError("Excel dosyasında 'text_data' sütunu bulunamadı")
            texts = df['text_data'].dropna().tolist()
        else:
            if input_path.lower().endswith(".pdf"):
                text_content = extract_text_from_pdf(input_path)
            else:
                with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                    text_content = f.read()
            # Stage 1: chunking
            texts = chunk_text_sliding_window(text_content)

        if not texts:
            raise ValueError("Analiz edilecek metin veya chunk bulunamadı.")

        total = len(texts)
        jobs[job_id]["total"] = total
        save_job(job_id)
        log_job(job_id, f"Analiz edilecek toplam segment sayisi: {total}", "system")
        
        pipeline = QualitativeAnalysisPipeline()
        
        sem = asyncio.Semaphore(2) # Maksimum 2 eszamanli islem
        processed_count = 0
        
        async def process_with_sem(t: str, idx: int):
            nonlocal processed_count
            if jobs[job_id]["status"] == "cancelled":
                raise asyncio.CancelledError()
            
            async with sem:
                await asyncio.sleep(3) # API limitleri icin akilli bekleme
                if jobs[job_id]["status"] == "cancelled":
                    raise asyncio.CancelledError()
                
                try:
                    res = await pipeline.process_single_text(t, idx, job_id)
                except asyncio.CancelledError:
                    raise asyncio.CancelledError()
                
                processed_count += 1
                jobs[job_id]["processed"] = processed_count
                jobs[job_id]["progress"] = int((processed_count / total) * 90)
                save_job(job_id)
                return res

        try:
            log_job(job_id, f"Semaphore tabanli akilli kuyruk baslatiliyor... (Maks {sem._value} eszamanli)", "system")
            tasks = [process_with_sem(t, i + 1) for i, t in enumerate(texts)]
            pipeline.results = await asyncio.gather(*tasks)
        except asyncio.CancelledError:
            log_job(job_id, "Analiz durduruldu, islem iptal ediliyor.", "system")
            return

        # Stage 5: Reduce Phase (Academic Synthesis)
        if jobs[job_id]["status"] == "cancelled":
            log_job(job_id, "Analiz durduruldu, işlem iptal ediliyor.", "system")
            return
            
        log_job(job_id, "Sentez Asamasi (Reduce Phase) baslatiliyor...", "system")
        jobs[job_id]["status"] = "reducing"
        save_job(job_id)
        
        if pipeline.results:
            pipeline.report, pipeline.report_model, pipeline.report_fallback_used = await pipeline.agent_d.synthesize_report(pipeline.results, job_id)
            
        pipeline.save_results(output_path)

        jobs[job_id]["results"] = [
            {
                "original_text": r.original_text,
                "pii_redacted_text": r.pii_redacted_text,
                "mapped_themes": r.mapped_themes,
                "validation_result": r.validation_result,
                "validation_reason": r.validation_reason,
                "theme_model": r.theme_model,
                "validation_model": r.validation_model,
                "theme_fallback_used": r.theme_fallback_used,
                "validation_fallback_used": r.validation_fallback_used,
                "pii_redaction_count": r.pii_redaction_count,
            }
            for r in pipeline.results
        ]
        
        jobs[job_id]["report"] = pipeline.report
        jobs[job_id]["report_model"] = pipeline.report_model
        jobs[job_id]["report_fallback_used"] = pipeline.report_fallback_used
        jobs[job_id]["privacy_metrics"] = {
            "total_redactions": sum(r.pii_redaction_count for r in pipeline.results),
            "segments_with_redactions": sum(1 for r in pipeline.results if r.pii_redaction_count > 0),
            "total_segments": len(pipeline.results),
        }
        jobs[job_id]["model_summary"] = {
            "theme_models": sorted({r.theme_model for r in pipeline.results if r.theme_model}),
            "validation_models": sorted({r.validation_model for r in pipeline.results if r.validation_model}),
            "report_model": pipeline.report_model,
            "fallback_count": sum(1 for r in pipeline.results if r.theme_fallback_used or r.validation_fallback_used) + (1 if pipeline.report_fallback_used else 0),
        }
        jobs[job_id]["status"] = "done"
        jobs[job_id]["progress"] = 100
        save_job(job_id)
        log_job(job_id, "Analiz basariyla tamamlandi. Rapor hazirlandı.", "system")

    except Exception as e:
        log_job(job_id, f"Is hatayla sonlandi: {str(e)}", "system", is_error=True)
        jobs[job_id]["status"] = "error"
        jobs[job_id]["error"] = str(e)
        save_job(job_id)
    finally:
        # Clean up uploaded input file
        try:
            os.remove(input_path)
        except Exception:
            pass

# ---------------------------------------------------------------------------
# API Endpoints
# ---------------------------------------------------------------------------

@app.post("/upload")
async def upload_file(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    """Upload Excel, TXT or PDF file and start analysis"""
    filename = file.filename.lower()
    if not filename.endswith(('.xlsx', '.xls', '.txt', '.pdf')):
        raise HTTPException(status_code=400, detail="Sadece .xlsx, .xls, .txt veya .pdf dosyaları kabul edilir")

    job_id = str(uuid.uuid4())
    suffix = Path(file.filename).suffix
    input_path = f"uploads/{job_id}_input{suffix}"
    output_path = f"outputs/{job_id}_output.xlsx"

    os.makedirs("uploads", exist_ok=True)
    os.makedirs("outputs", exist_ok=True)

    content = await file.read()
    with open(input_path, "wb") as f:
        f.write(content)

    is_excel = filename.endswith(('.xlsx', '.xls'))

    jobs[job_id] = {
        "status": "queued",
        "progress": 0,
        "processed": 0,
        "total": 0,
        "output_path": output_path,
        "results": [],
        "report": "",
        "error": None,
        "logs": [],
        "router_events": [],
        "privacy_metrics": {"total_redactions": 0, "segments_with_redactions": 0, "total_segments": 0},
        "model_summary": {},
        "report_model": "",
        "report_fallback_used": False,
        "filename": file.filename,
        "created_at": datetime.now().isoformat()
    }
    save_job(job_id)

    background_tasks.add_task(run_pipeline_job, job_id, input_path, output_path, is_excel)
    return {"job_id": job_id}


@app.get("/jobs")
async def get_jobs():
    """List all jobs in database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT job_id, filename, status, progress, total, created_at FROM jobs ORDER BY created_at DESC")
    rows = cursor.fetchall()
    conn.close()
    
    return [
        {
            "job_id": row[0],
            "filename": row[1],
            "status": row[2],
            "progress": row[3],
            "total": row[4],
            "created_at": row[5]
        }
        for row in rows
    ]


@app.get("/status/{job_id}")
async def get_status(job_id: str):
    """Get job status and progress"""
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job bulunamadı")
    job = jobs[job_id]
    return {
        "status": job["status"],
        "progress": job["progress"],
        "processed": job["processed"],
        "total": job["total"],
        "error": job["error"],
        "report": job.get("report", ""),
        "logs": job.get("logs", []),
        "router_events": job.get("router_events", []),
        "privacy_metrics": job.get("privacy_metrics", {}),
        "model_summary": job.get("model_summary", {}),
    }


@app.get("/results/{job_id}")
async def get_results(job_id: str):
    """Get analysis results as JSON"""
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job bulunamadı")
    if jobs[job_id]["status"] != "done":
        raise HTTPException(status_code=400, detail="Analiz henüz tamamlanmadı")
    return {
        "results": jobs[job_id]["results"],
        "report": jobs[job_id]["report"],
        "router_events": jobs[job_id].get("router_events", []),
        "privacy_metrics": jobs[job_id].get("privacy_metrics", {}),
        "model_summary": jobs[job_id].get("model_summary", {}),
        "report_model": jobs[job_id].get("report_model", ""),
        "report_fallback_used": jobs[job_id].get("report_fallback_used", False),
    }


@app.get("/download/{job_id}")
async def download_results(job_id: str):
    """Download results as Excel file"""
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job bulunamadı")
    if jobs[job_id]["status"] != "done":
        raise HTTPException(status_code=400, detail="Analiz henüz tamamlanmadı")
    output_path = jobs[job_id]["output_path"]
    if not Path(output_path).exists():
        raise HTTPException(status_code=404, detail="Çıktı dosyası bulunamadı")
    return FileResponse(
        path=output_path,
        filename="analysis_results.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@app.get("/download/report/{job_id}")
async def download_report(job_id: str):
    """Download report as Markdown file"""
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job bulunamadı")
    if jobs[job_id]["status"] != "done":
        raise HTTPException(status_code=400, detail="Analiz henüz tamamlanmadı")
    output_path = jobs[job_id]["output_path"]
    report_path = output_path.replace("_output.xlsx", "_report.md").replace(".xlsx", "_report.md")
    if not Path(report_path).exists():
        raise HTTPException(status_code=404, detail="Rapor dosyası bulunamadı")
    return FileResponse(
        path=report_path,
        filename="academic_report.md",
        media_type="text/markdown"
    )


@app.get("/download/report/docx/{job_id}")
async def download_report_docx(job_id: str):
    """Generate and download report as Word (.docx) file"""
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job bulunamadı")
    job = jobs[job_id]
    if job["status"] != "done":
        raise HTTPException(status_code=400, detail="Analiz henüz tamamlanmadı")
    
    report_text = job.get("report", "")
    if not report_text:
        raise HTTPException(status_code=404, detail="Rapor içeriği bulunamadı")
        
    # Generate Word Document
    from docx import Document
    from docx.shared import Pt
    
    doc = Document()
    
    # Title formatting
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Med-AgentLab Nitel Analiz Akademik Raporu")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_p.alignment = 1 # Center
    
    # Add meta information
    privacy = job.get("privacy_metrics", {})
    model_summary = job.get("model_summary", {})
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Dosya Adı: {job.get('filename', 'Bilinmeyen Dosya')}\n")
    meta_p.add_run(f"Analiz Tarihi: {job.get('created_at', '')[:10]}\n")
    meta_p.add_run(f"Segment Sayısı: {job.get('total', 0)}\n")
    meta_p.add_run(f"PII Redaksiyon Sayısı: {privacy.get('total_redactions', 0)}\n")
    meta_p.add_run(f"Rapor Modeli: {job.get('report_model') or model_summary.get('report_model') or 'Bilinmiyor'}\n")
    meta_p.add_run(f"Fallback Kullanımı: {'Evet' if model_summary.get('fallback_count', 0) else 'Hayır'}\n")
    meta_p.paragraph_format.space_after = Pt(24)

    pipeline_p = doc.add_paragraph()
    pipeline_title = pipeline_p.add_run("Pipeline Özeti")
    pipeline_title.bold = True
    pipeline_title.font.size = Pt(14)
    for item in [
        "1. Excel/TXT/PDF girdisi metin segmentlerine dönüştürülür.",
        "2. Yerel Ollama + regex katmanı kişisel verileri maskeler.",
        "3. Tema haritalama görevi kota-duyarlı model havuzuyla çalışır.",
        "4. PubMed destekli validasyon modeli temaları denetler.",
        "5. Reducer havuzu doğrulanmış temalardan akademik rapor üretir.",
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    # Basic Markdown parsing to Word paragraph structures
    lines = report_text.split("\n")
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
            
        if line_strip.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line_strip[2:])
            r.bold = True
            r.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line_strip.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line_strip[3:])
            r.bold = True
            r.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
        elif line_strip.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line_strip[4:])
            r.bold = True
            r.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif line_strip.startswith("- ") or line_strip.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            # Handle inline bold tags (**text**)
            parts = re.split(r'(\*\*.*?\*\*)', line_strip[2:])
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line_strip)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
            p.paragraph_format.space_after = Pt(6)
            
    # Save document to file
    docx_path = f"outputs/{job_id}_report.docx"
    doc.save(docx_path)
    
    return FileResponse(
        path=docx_path,
        filename="academic_report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.post("/cancel/{job_id}")
async def cancel_job(job_id: str):
    """Cancel a running analysis job"""
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job bulunamadı")
    
    current_status = jobs[job_id]["status"]
    if current_status in ["queued", "running", "reducing"]:
        jobs[job_id]["status"] = "cancelled"
        log_job(job_id, "Analiz kullanıcı tarafından iptal edildi.", "system", is_error=True)
        save_job(job_id)
        return {"status": "cancelled"}
    return {"status": current_status, "message": "İş aktif bir durumda değil (zaten bitmiş veya iptal edilmiş olabilir)"}


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.get("/")
async def serve_frontend():
    """Serve the frontend SPA"""
    return FileResponse("frontend/index.html")
